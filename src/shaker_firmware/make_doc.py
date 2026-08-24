# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ---- 한글 폰트(맑은 고딕) 기본 설정 ----
def set_kr_font(doc, name="맑은 고딕", size=11):
    st = doc.styles["Normal"]
    st.font.name = name
    st.font.size = Pt(size)
    rpr = st.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)

set_kr_font(doc)

def mono(text):
    p = doc.add_paragraph()
    for ln in text.split("\n"):
        r = p.add_run(ln + "\n")
        r.font.name = "Consolas"
        r.font.size = Pt(9.5)
        r._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Consolas")
    return p

# ===== Title =====
h = doc.add_heading("", level=0)
run = h.add_run("이중 서보 조화가진장치 — IMU 기반 위상동기 제어 구조")
run.font.size = Pt(18)
doc.add_paragraph("작성: 제어구조 정리 (블록다이어그램 작도용)").italic = True

# 1
doc.add_heading("1. 시스템 개요 & 목표", level=1)
doc.add_paragraph("연속회전 서보 2개가 크랭크-슬라이더로 30×30 cm 판을 상하 가진(≈ 2 Hz)한다.")
p = doc.add_paragraph(); p.add_run("두 가지 목표를 동시에 달성:").bold = True
doc.add_paragraph("수직 조화가진 z̈ 생성 (장치의 목적 기능)", style="List Number")
doc.add_paragraph("판의 롤 θ(기울기)를 0으로 유지 (두 서보 위상 드리프트로 판이 점점 기울어지는 문제 보정)", style="List Number")

# 2
doc.add_heading("2. 핵심 물리 — 두 운동이 분리되어 한 구조에 공존", level=1)
doc.add_paragraph("두 크랭크 각도 φ₁, φ₂에 대해 판의 두 운동은 공통모드/차동모드로 분리된다:")
t = doc.add_table(rows=1, cols=3); t.style = "Light Grid Accent 1"
hdr = t.rows[0].cells
hdr[0].text = "운동"; hdr[1].text = "결정 요인"; hdr[2].text = "성격"
rows = [
    ("수직 가진 z", "φ₁ + φ₂ (공통/합)", "만들고 싶은 출력 → u₀로 구동 (개루프)"),
    ("롤 θ", "φ₁ − φ₂ (차동/차)", "0으로 없애고 싶은 것 → 피드백으로 규제"),
]
for a, b, c in rows:
    cells = t.add_row().cells
    cells[0].text = a; cells[1].text = b; cells[2].text = c
doc.add_paragraph("핵심: 위상차(Δ)를 조절해 롤을 잡아도 두 크랭크의 평균 회전(=수직 가진)은 그대로 유지된다 → 서로 간섭 없이 한 구조에 공존한다.")

# 3
doc.add_heading("3. 플랜트 (입력 2 / 출력 2)", level=1)
doc.add_paragraph("플랜트 = [서보 2개 + 크랭크-슬라이더 + 판] 한 블록")
p = doc.add_paragraph(); p.add_run("입력:").bold = True
doc.add_paragraph("u₀ (BASE_OFFSET): 두 서보 공통 기준 구동속도. 가진을 만드는 구동력 (외부 설정값, 조작변수 아님).", style="List Bullet")
doc.add_paragraph("Δ: 슬레이브 서보 속도 보정 = 유일한 조작변수. 마스터=u₀ 고정, 슬레이브 명령=u₀+Δ.", style="List Bullet")
p = doc.add_paragraph(); p.add_run("출력:").bold = True
doc.add_paragraph("z̈: 수직 조화가진(≈2 Hz) — 목적 출력, 피드백 안 함 (개루프).", style="List Bullet")
doc.add_paragraph("θ: 판의 롤 — 규제 대상, 피드백 함 (폐루프).", style="List Bullet")

# 4
doc.add_heading("4. 순방향 경로 (forward path)", level=1)
steps = [
    "기준입력: θ_ref = 0 (판 수평이 목표)",
    "비교기 (합산점 ⊗): e = θ_ref − θ_meas. 피드백은 음(−)으로 입력.",
    "PI 제어기 (병렬 2블록 → 합산점에서 합): 비례항 K_p·e, 적분항 K_i∫e dt (위상 드리프트 제거의 핵심), 합산점(+/+) → 보정량 Δ 출력.",
    "Δ → 플랜트 (슬레이브 속도 보정으로 인가).",
    "플랜트가 z̈(가진 출력)와 θ(롤)를 동시에 생성.",
]
for s in steps:
    doc.add_paragraph(s, style="List Number")

# 5
doc.add_heading("5. 피드백 경로 (feedback path) — 상관검파로 롤/위상오차 측정", level=1)
fb = [
    "IMU (MPU-6500): 자이로 ω_roll(롤 각속도) + 가속도계 a_z(수직 가속도) 측정.",
    "High-pass 필터: 중력/DC 제거 → 교류 성분 ω̃_roll, ã_z.",
    "곱셈기 (×): ω̃_roll · ã_z (순간 상관).",
    "Low-pass 필터 ⟨·⟩: 시간 평균 → ⟨ω̃_roll · ã_z⟩ ∝ sin(Δφ) ≈ 위상오차/롤.",
    "이 값을 비교기(−)로 환류 → 루프 폐합.",
]
for s in fb:
    doc.add_paragraph(s, style="List Number")
p = doc.add_paragraph()
p.add_run("원리: ").bold = True
p.add_run("두 크랭크가 동기되면 판이 수평 유지하며 상하로만 움직여 롤≈0. 어긋나면 롤이 생기고 ⟨롤각속도×수직가속도⟩의 평균이 sin(위상차)에 비례 → 드라이브 주파수를 몰라도 오차의 크기·부호 추출(동기검파). PI 적분기가 이 오차를 0으로 몰아 슬레이브 평균속도를 마스터에 맞춘다.")

# 6
doc.add_heading("6. 블록 배치 (그릴 때 순서)", level=1)
doc.add_paragraph("순방향 (위, 좌→우):")
mono(
"[theta_ref=0] -> (+/-)O -> e -> [ K_p ]  \\\n"
"                              (+/+)O -> Delta -> [ Plant: 2서보+크랭크+판 ] -> z.. 가진출력(~2Hz)\n"
"                     e -> [ K_i∫dt ]  /                ^                         |\n"
"                                               u0(BASE_OFFSET)                  theta(롤)"
)
doc.add_paragraph("피드백 (아래, 우→좌):")
mono(
"theta -> [IMU: w_roll, a_z] -> [High-pass] -> [ x 곱셈 ] -> [Low-pass <.>] -> (-)O 로 환류"
)
doc.add_paragraph("블록 색 구분(권장): 흰색=기준입력, 주황=게인(K_p,K_i), 남색=처리(High-pass/×/Low-pass), 초록=센서(IMU), 하늘색=플랜트.")
doc.add_paragraph("합산점 vs 곱셈기: 합산점 ⊗는 +/− 부호 표시, 곱셈기는 × 로 구분.")

# 7
doc.add_heading("7. 실제 파라미터 (참고)", level=1)
t2 = doc.add_table(rows=1, cols=2); t2.style = "Light Grid Accent 1"
t2.rows[0].cells[0].text = "항목"; t2.rows[0].cells[1].text = "값"
params = [
    ("기준속도 u₀", "400 (정지펄스 1513 µs 기준 오프셋), 가진 ≈ 2 Hz"),
    ("PI 게인", "K_p = 8, K_i = 40 (적분 위주 설계)"),
    ("보정 범위 Δ", "±250 µs"),
    ("IMU 축", "수직가속 = accel Z, 롤레이트 = gyro Y (실측 확정)"),
    ("센서", "MPU-6500 (I²C, 100 kHz)"),
    ("서보 정지펄스", "1513 µs (DS3345 데드밴드 1504~1522 중앙)"),
]
for a, b in params:
    c = t2.add_row().cells; c[0].text = a; c[1].text = b

out = r"c:\Users\chang\OneDrive\문서\PlatformIO\Projects\servo_test\control_structure.docx"
doc.save(out)
print("saved", out)
